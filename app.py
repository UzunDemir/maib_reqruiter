import os
import streamlit as st
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import time
import json
from PyPDF2 import PdfReader
import tempfile
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import docx
from docx import Document
import io
from concurrent.futures import ThreadPoolExecutor, as_completed

st.set_page_config(layout="wide", initial_sidebar_state="expanded")


# --- Language Settings ---
if 'language' not in st.session_state:
    st.session_state.language = 'rom'  # Default language

def get_translation(key):
    translations = {
        'rom': {
            'app_title': 'AI HR-Recruiter',
            'sidebar_title': 'MAIB AI HR-Recruiter',
            'load_vacancies': 'Încarcă ofertele de muncă pentru tine...',
            'vacancies_list': 'Lista ofertelor MAIB:',
            'upload_cv': 'Încărcă CV-ul tău (PDF, DOCX, TXT)',
            'best_matches': 'Cele mai relevante oferte pentru CV-ul tău',
            'generate_analysis': 'Generează analiza de potrivire',
            'detailed_analysis': 'Analiză detaliată a conformității',
            'download_analysis': 'Descarcă analiza (DOCX)',
            'start_interview': 'A trece interviul introductiv',
            'interview_in_progress': 'Interviul a început! Vă rog să răspundeți la întrebările de mai jos.',
            'finish_interview': 'Interviul s-a încheiat',
            'start_technical': 'Începe interviul tehnic',
            'technical_in_progress': 'Interviul tehnic a început! Vă rugăm să răspundeți la întrebările de mai jos.',
            'finish_technical': 'Finalizează interviul tehnic',
            'technical_feedback': 'Feedback tehnic',
            'final_conclusion': 'Concluzia finală',
            'reset_process': 'Resetează procesul',
            'candidate_profile': 'Profilul candidatului',
            'download_profile': 'Descarcă profilul candidatului (DOCX)',
            'current_step': 'Etapa curentă:',
            'step1': 'Încărcare CV',
            'step2': 'Analiză potrivire',
            'step3': 'Interviu general',
            'step4': 'Interviu tehnic',
            'step5': 'Concluzie finală'
        },
        'rus': {
            'app_title': 'AI HR-Рекрутер',
            'sidebar_title': 'MAIB AI HR-Рекрутер',
            'load_vacancies': 'Загрузить вакансии для вас...',
            'vacancies_list': 'Список вакансий MAIB:',
            'upload_cv': 'Загрузите ваше резюме (PDF, DOCX, TXT)',
            'best_matches': 'Самые подходящие вакансии для вашего резюме',
            'generate_analysis': 'Сгенерировать анализ соответствия',
            'detailed_analysis': 'Подробный анализ соответствия',
            'download_analysis': 'Скачать анализ (DOCX)',
            'start_interview': 'Начать общее собеседование',
            'interview_in_progress': 'Собеседование начато! Пожалуйста, ответьте на вопросы ниже.',
            'finish_interview': 'Завершить собеседование',
            'start_technical': 'Начать техническое собеседование',
            'technical_in_progress': 'Техническое собеседование начато! Пожалуйста, ответьте на вопросы ниже.',
            'finish_technical': 'Завершить техническое собеседование',
            'technical_feedback': 'Технический отзыв',
            'final_conclusion': 'Заключительный вывод',
            'reset_process': 'Сбросить процесс',
            'candidate_profile': 'Профиль кандидата',
            'download_profile': 'Скачать профиль кандидата (DOCX)',
            'current_step': 'Текущий этап:',
            'step1': 'Загрузка резюме',
            'step2': 'Анализ соответствия',
            'step3': 'Общее собеседование',
            'step4': 'Техническое собеседование',
            'step5': 'Заключительный вывод'
        },
        'en': {
            'app_title': 'AI HR-Recruiter',
            'sidebar_title': 'MAIB AI HR-Recruiter',
            'load_vacancies': 'Load job vacancies for you...',
            'vacancies_list': 'List of MAIB vacancies:',
            'upload_cv': 'Upload your CV (PDF, DOCX, TXT)',
            'best_matches': 'Most relevant job matches for your CV',
            'generate_analysis': 'Generate matching analysis',
            'detailed_analysis': 'Detailed compatibility analysis',
            'download_analysis': 'Download analysis (DOCX)',
            'start_interview': 'Start introductory interview',
            'interview_in_progress': 'Interview has started! Please answer the questions below.',
            'finish_interview': 'Interview completed',
            'start_technical': 'Start technical interview',
            'technical_in_progress': 'Technical interview has started! Please answer the questions below.',
            'finish_technical': 'Complete technical interview',
            'technical_feedback': 'Technical feedback',
            'final_conclusion': 'Final conclusion',
            'reset_process': 'Reset process',
            'candidate_profile': 'Candidate profile',
            'download_profile': 'Download candidate profile (DOCX)',
            'current_step': 'Current step:',
            'step1': 'CV Upload',
            'step2': 'Matching Analysis',
            'step3': 'General Interview',
            'step4': 'Technical Interview',
            'step5': 'Final Conclusion'
        }
    }
    lang = st.session_state.get('language', 'rom')
    return translations.get(lang, translations['rom']).get(key, key)


# --- Stiluri și bara laterală ---
st.markdown(f"""
    <style>
        /* 1. ПОЛНОЕ СКРЫТИЕ ХЕДЕРА (кроме стрелки сайдбара) */
        header div:nth-child(2) {{
            display: none !important;
        }}
        [data-testid="stHeader"] {{
            background: rgba(0,0,0,0) !important;
        }}
        [data-testid="stSidebarCollapseButton"] {{
            visibility: visible !important;
        }}
        /* Цвет стрелки */
        [data-testid="stSidebarCollapseButton"] svg {{
            fill: white !important;
            color: white !important;
        }}
        
        /* Убираем лишний отступ сверху */
        .block-container {{
            padding-top: 1rem !important;
        }}

        /* 2. ВАШИ СТИЛИ */
        section[data-testid="stSidebar"] {{
            background-color: #253646 !important;
        }}
        .sidebar-title {{
            color: white;
            font-size: 24px;
            font-weight: bold;
            text-align: center;
            margin-bottom: 1rem;
        }}
        .sidebar-text {{
            color: white;
        }}
        #MainMenu, footer {{
            display: none !important;
        }}
        .center {{
            display: flex;
            justify-content: center;
            align-items: center;
            text-align: center;
            flex-direction: column;
            margin-top: 0vh;
        }}
        /* ... остальные ваши стили (match-card, progress-bar и т.д.) ... */
        .match-card {{ border-radius: 10px; padding: 15px; margin-bottom: 15px; background-color: #f0f2f6; }}
        .match-header {{ display: flex; justify-content: space-between; align-items: center; }}
        .progress-bar {{ height: 10px; background-color: #e0e0e0; border-radius: 5px; margin-top: 5px; }}
        .progress-fill {{ height: 100%; border-radius: 5px; background-color: #40c1ac; }}
        .current-step {{ background-color: #40c1ac; color: white; padding: 10px; border-radius: 5px; margin-top: 10px; font-weight: bold; }}
    </style>
    <div class="center">
        <img src="https://www.maib.md/images/logo.svg" width="300">
        <h1>{get_translation('app_title')}</h1>
    </div>
""", unsafe_allow_html=True)



# --- Sidebar Content ---
with st.sidebar:
    st.markdown('<div class="sidebar-title">{}</div>'.format(get_translation('sidebar_title')), unsafe_allow_html=True)
    
    # Language selector
    language = st.radio("Language / Язык / Limbă:", 
                       ["rom", "rus", "en"],
                       index=["rom", "rus", "en"].index(st.session_state.language),
                       key="lang_selector",
                       horizontal=True)
    
    if language != st.session_state.language:
        st.session_state.language = language
        st.rerun()
    
    st.divider()
    
    # Current step indicator
    current_step = 1
    if 'knowledge_base' in st.session_state and st.session_state.knowledge_base.uploaded_files:
        current_step = 2
    if 'analysis' in st.session_state and st.session_state.analysis:
        current_step = 3
    if 'profile' in st.session_state and st.session_state.profile:
        current_step = 4
    if 'final_recommendation' in st.session_state and st.session_state.final_recommendation:
        current_step = 5

        
    st.markdown(f'<div class="current-step">{get_translation("current_step")} {current_step}/5</div>', unsafe_allow_html=True)
    st.markdown(f"- {get_translation(f'step{current_step}')}")
    
    st.divider()
    
    # Process description
    st.markdown("""
    <div class="sidebar-text">
    
    1. 📥 <strong>{}</strong>  
   <em>{}</em>

    2. 📄 <strong>{}</strong>  
       <em>{}</em>

    3. 🤖 <strong>{}</strong>  
       <em>{}</em>

    4. 🔍 <strong>{}</strong>  
       <em>{}<br>{}</em>

    5. ✅ <strong>{}</strong>  
       <em>{}</em>

    6. 🗣️ <strong>{}</strong>  
       <em>{}</em>

    7. ⚡ <strong>{}</strong>    
       <em>{}</em>

    8. 💻 <strong>{}</strong>  
       <em>{}</em>

    9. 📋 <strong>{}</strong>  
       <em>{}</em>
    </div>
    """.format(
        get_translation('step1'), "Agentul încarcă automat toate posturile vacante actuale de la MAIB." if st.session_state.language == 'rom' else 
        "Агент автоматически загружает все текущие вакансии MAIB." if st.session_state.language == 'rus' else 
        "The agent automatically loads all current MAIB vacancies.",
        
        get_translation('step2'), "Utilizatorul își încarcă CV-ul pentru analiză." if st.session_state.language == 'rom' else 
        "Пользователь загружает свое резюме для анализа." if st.session_state.language == 'rus' else 
        "The user uploads their CV for analysis.",
        
        get_translation('step3'), "Agentul analizează CV-ul și identifică top 3 posturi relevante." if st.session_state.language == 'rom' else 
        "Агент анализирует резюме и определяет топ-3 подходящих вакансий." if st.session_state.language == 'rus' else 
        "The agent analyzes the CV and identifies top 3 relevant positions.",
        
        get_translation('step4'), "Evidențiază punctele forte ale candidatului." if st.session_state.language == 'rom' else 
        "Выделяет сильные стороны кандидата." if st.session_state.language == 'rus' else 
        "Highlights the candidate's strengths.",
        
        "Identifică punctele slabe sau lipsurile în competențe." if st.session_state.language == 'rom' else 
        "Определяет слабые стороны или пробелы в навыках." if st.session_state.language == 'rus' else 
        "Identifies weaknesses or skill gaps.",
        
        get_translation('step5'), "Dacă este interesat, candidatul își exprimă acordul pentru a continua procesul." if st.session_state.language == 'rom' else 
        "Если заинтересован, кандидат выражает согласие продолжить процесс." if st.session_state.language == 'rus' else 
        "If interested, the candidate agrees to continue the process.",
        
        get_translation('step3'), "Agentul pune întrebări generale, analizează răspunsurile și formulează primele concluzii." if st.session_state.language == 'rom' else 
        "Агент задает общие вопросы, анализирует ответы и формулирует первые выводы." if st.session_state.language == 'rus' else 
        "The agent asks general questions, analyzes responses and formulates initial conclusions.",
        
        get_translation('step4'), "Identificarea și verificarea textelor create automat pentru a asigura autenticitatea conținutului." if st.session_state.language == 'rom' else 
        "Выявление и проверка автоматически созданных текстов для обеспечения подлинности." if st.session_state.language == 'rus' else 
        "Identifying and verifying automatically generated text to ensure content authenticity.",
        
        get_translation('step4'), "Evaluarea competențelor tehnice și furnizarea unui feedback tehnic." if st.session_state.language == 'rom' else 
        "Оценка технических навыков и предоставление технического отзыва." if st.session_state.language == 'rus' else 
        "Assessing technical skills and providing technical feedback.",
        
        get_translation('step5'), "Agentul oferă un verdict final: recomandare pentru angajare sau refuz argumentat." if st.session_state.language == 'rom' else 
        "Агент выносит окончательный вердикт: рекомендация к найму или обоснованный отказ." if st.session_state.language == 'rus' else 
        "The agent provides a final verdict: hiring recommendation or justified refusal."
    ), unsafe_allow_html=True)

    
    # Vacancies list in alphabetical order
    if 'vacancies_data' in st.session_state and st.session_state.vacancies_data:
        st.divider()
        st.markdown(f"### 🔎 {get_translation('vacancies_list')}")
        st.success(f"Oferte găsite: {len(st.session_state.vacancies_data)}" if st.session_state.language == 'rom' else 
                 f"Найдено вакансий: {len(st.session_state.vacancies_data)}" if st.session_state.language == 'rus' else 
                 f"Found vacancies: {len(st.session_state.vacancies_data)}")
        
        # Sort vacancies alphabetically
        sorted_vacancies = sorted(st.session_state.vacancies_data, key=lambda x: x['title'])
        for vac in sorted_vacancies:
            st.markdown(
                f'<a href="{vac["url"]}" target="_blank" style="color:#40c1ac; text-decoration:none;">• {vac["title"]}</a>',
                unsafe_allow_html=True
            )

st.divider()

#######################################################################
class DocumentChunk:
    def __init__(self, text, doc_name, page_num):
        self.text = text
        self.doc_name = doc_name
        self.page_num = page_num

class KnowledgeBase:
    def __init__(self):
        self.chunks = []
        self.uploaded_files = []
        self.doc_texts = []
    
    def clear(self):
        self.chunks = []
        self.doc_texts = []
        self.uploaded_files = []
    
    def split_text(self, text, max_chars=2000):
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + max_chars, len(text))
            chunk = text[start:end].strip()
            if chunk:  # Skip empty chunks
                chunks.append(chunk)
            start = end
        return chunks

    def load_pdf(self, file_content, file_name):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name

            with open(tmp_file_path, 'rb') as file:
                reader = PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        for chunk in self.split_text(page_text):
                            self.chunks.append(DocumentChunk(chunk, file_name, page_num+1))
                            self.doc_texts.append(chunk)
            self.uploaded_files.append(file_name)
            return True
        except Exception as e:
            st.error(f"Eroare la încărcarea PDF: {str(e)}")
            return False
        finally:
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)

    def load_docx(self, file_content, file_name):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name

            doc = docx.Document(tmp_file_path)
            full_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    full_text.append(text)
            text = "\n".join(full_text)
            for chunk in self.split_text(text):
                self.chunks.append(DocumentChunk(chunk, file_name, 0))
                self.doc_texts.append(chunk)
            self.uploaded_files.append(file_name)
            return True
        except Exception as e:
            st.error(f"Eroare la încărcarea DOCX: {str(e)}")
            return False
        finally:
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)

    def load_txt(self, file_content, file_name):
        try:
            text = file_content.decode('utf-8', errors='ignore')
            for chunk in self.split_text(text):
                self.chunks.append(DocumentChunk(chunk, file_name, 0))
                self.doc_texts.append(chunk)
            self.uploaded_files.append(file_name)
            return True
        except Exception as e:
            st.error(f"Eroare la încărcarea TXT: {str(e)}")
            return False

    def load_file(self, uploaded_file):
        name = uploaded_file.name.lower()
        content = uploaded_file.read()
        if name.endswith('.pdf'):
            return self.load_pdf(content, uploaded_file.name)
        elif name.endswith('.docx'):
            return self.load_docx(content, uploaded_file.name)
        elif name.endswith('.txt'):
            return self.load_txt(content, uploaded_file.name)
        else:
            st.warning(f"Formatul fișierului {uploaded_file.name} nu este suportat.")
            return False

    def get_all_text(self):
        return "\n\n".join(self.doc_texts)
#########################################################

# --- Initialize knowledge base ---
if 'knowledge_base' not in st.session_state:
    st.session_state.knowledge_base = KnowledgeBase()

if 'vacancies_data' not in st.session_state:
    st.session_state.vacancies_data = []



##############################
# Încărcare oferte de pe rabota.md
def scrape_vacancy(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        soup_vac = BeautifulSoup(resp.text, 'html.parser')

        title_tag = soup_vac.find('h1')
        title = title_tag.get_text(strip=True) if title_tag else 'Titlu indisponibil'

        vacancy_content = soup_vac.find('div', class_='vacancy-content')
        description = vacancy_content.get_text(separator='\n', strip=True) if vacancy_content else 'Descriere indisponibilă'

        return {'url': url, 'title': title, 'description': description}
    except Exception as e:
        st.error(f"Eroare la preluarea ofertei {url}: {str(e)}")
        return None

#######################################################################################################################
from concurrent.futures import ThreadPoolExecutor, as_completed

def load_vacancies():
    base_url = "https://www.rabota.md/ru/companies/moldova-agroindbank#vacancies"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    with st.spinner("Se încarcă ofertele de muncă..."):
        try:
            response = requests.get(base_url, headers=headers, timeout=10)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')

            links = soup.find_all('a', class_='vacancyShowPopup')
            urls = [urljoin(base_url, a['href']) for a in links]

            vacancies_data = []
            progress_bar = st.progress(0)
            status_text = st.empty()

            total = len(urls)
            done = 0

            with ThreadPoolExecutor(max_workers=5) as executor:
                future_to_url = {executor.submit(scrape_vacancy, url): url for url in urls}
                
                for future in as_completed(future_to_url):
                    result = future.result()
                    if result:
                        vacancies_data.append(result)
                        done += 1
                        progress_bar.progress(done / total)
                        status_text.text(f"[{done}/{total}] Ofertă încărcată: {result['title']}")

            st.session_state.vacancies_data = vacancies_data
            st.success(f"Au fost încărcate {len(vacancies_data)} oferte de muncă!")

        except Exception as e:
            st.error(f"Eroare la încărcarea ofertelor: {str(e)}")
##########################################################################################################################################

if st.button(get_translation('load_vacancies')):
    load_vacancies()
    st.rerun()

# --- CV Upload Section ---
st.markdown(f"### 📄 {get_translation('upload_cv')}")
uploaded_files = st.file_uploader(get_translation('upload_cv'), type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
####################################################################################

####################################################################################
if uploaded_files:
    kb = st.session_state.knowledge_base
    kb.clear()
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, uploaded_file in enumerate(uploaded_files):
        success = kb.load_file(uploaded_file)
        progress = (i + 1) / len(uploaded_files)
        progress_bar.progress(progress)
        status_text.text(f"Se procesează fișierul {i+1}/{len(uploaded_files)}: {uploaded_file.name}" if st.session_state.language == 'rom' else
                        f"Обрабатывается файл {i+1}/{len(uploaded_files)}: {uploaded_file.name}" if st.session_state.language == 'rus' else
                        f"Processing file {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
    
    st.session_state.knowledge_base = kb
    st.success(f"Au fost încărcate {len(uploaded_files)} fișiere!" if st.session_state.language == 'rom' else
              f"Загружено файлов: {len(uploaded_files)}" if st.session_state.language == 'rus' else
              f"Uploaded files: {len(uploaded_files)}")

if not st.session_state.knowledge_base.uploaded_files:
    st.info("Te rugăm să încarci un CV pentru analiză" if st.session_state.language == 'rom' else
           "Пожалуйста, загрузите резюме для анализа" if st.session_state.language == 'rus' else
           "Please upload a CV for analysis")
    st.stop()

# --- Best Matches Section ---
st.markdown(f"### 🔍 {get_translation('best_matches')}")

cv_text = st.session_state.knowledge_base.get_all_text()
vacancies = st.session_state.vacancies_data

if not vacancies:
    st.warning("Nu există oferte de muncă disponibile. Te rugăm să încarci ofertele mai întâi." if st.session_state.language == 'rom' else
              "Нет доступных вакансий. Пожалуйста, сначала загрузите вакансии." if st.session_state.language == 'rus' else
              "No job vacancies available. Please load vacancies first.")
    st.stop()

# Analiza potrivirilor
#st.markdown("### 🔍 Cele mai relevante oferte pentru CV-ul tău")

cv_text = st.session_state.knowledge_base.get_all_text()
vacancies = st.session_state.vacancies_data

if not vacancies:
    st.warning("Nu există oferte de muncă disponibile. Te rugăm să încarci ofertele mai întâi.")
    st.stop()

# Procesare avansată a textelor
vacancy_texts = [f"{vac['title']}\n{vac['description']}" for vac in vacancies]
documents = [cv_text] + vacancy_texts

# TF-IDF
vectorizer = TfidfVectorizer(
    stop_words=None,
    ngram_range=(1, 2),  # Include bigrame pentru mai mult context
    max_features=5000
)

##############################################################################
################################
def check_if_ai_generated(answer_text):
    prompt = f"""
    Evaluează dacă următorul răspuns a fost generat de un om sau de o rețea neuronală (cum ar fi ChatGPT).

    Răspuns:
    \"\"\"
    {answer_text}
    \"\"\"

    Răspunde doar cu una dintre opțiuni:
    - uman
    - AI
    """

    response = requests.post(
        url,
        headers=headers,
        json={
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2
        }
    )

    result = response.json()['choices'][0]['message']['content'].strip().lower()
    return result

###########################################################




# --- Анализ совпадений вакансий ---
try:
    with st.spinner("Se analizează potrivirile..."):
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform(documents)
        similarity_scores = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()

    score_min, score_max = similarity_scores.min(), similarity_scores.max()
    if score_max - score_min > 0:
        normalized_scores = (similarity_scores - score_min) / (score_max - score_min) * 100
    else:
        normalized_scores = np.zeros_like(similarity_scores)

    normalized_scores = np.clip(normalized_scores, 0, 100)

    top_indices = similarity_scores.argsort()[::-1][:3]

    # Сохраняем top_indices в сессию, чтобы использовать дальше
    st.session_state.top_indices = top_indices

    for idx in top_indices:
        vac = vacancies[idx]
        score = normalized_scores[idx]

        with st.container():
            st.markdown(f"""
            <div class="match-card">
                <div class="match-header">
                    <h3><a href="{vac['url']}" target="_blank" style="text-decoration:none; color:inherit;">{vac['title']}</a></h3>
                    <h4>{score:.0f}% potrivire</h4>
                </div>
                <div class="progress-bar">
                    <div class="progress-fill" style="width: {score}%"></div>
                </div>
            </div>
            """, unsafe_allow_html=True)

        st.write("---")

except Exception as e:
    st.error(f"Eroare la analiza potrivirilor: {str(e)}")

# --- API Configuration ---
api_key = st.secrets.get("DEEPSEEK_API_KEY")
if not api_key:
    st.error("API ключ не настроен. Пожалуйста, добавьте его в Secrets.")
    st.stop()

url = "https://api.deepseek.com/v1/chat/completions"
headers = {
    "Authorization": f"Bearer {api_key}",
    "Content-Type": "application/json"
}


# Используем top_indices из сессии, если есть
top_indices = st.session_state.get("top_indices", [])
if len(top_indices) > 0:
    best_match_idx = top_indices[0]
    best_match_vacancy = vacancies[best_match_idx]

    # Отображаем кнопку для генерации анализа
    if st.button("🔍 Generează analiza de potrivire"):
        prompt = f"""
        Analizează corespondența dintre CV-ul candidatului și oferta de muncă.
        Mai întâi voi furniza CV-ul, apoi descrierea postului.

        CV-ul candidatului:
        {cv_text}
        
        Descrierea postului:
        {best_match_vacancy['description']}
        
        Vă rog să efectuați analiza conform următoarei structuri:

        1. Punctele forte ale CV-ului (potrivirea exactă cu cerințele postului)
        2. Punctele slabe sau lacunele din CV (unde candidatul nu corespunde)
        3. Recomandări concrete pentru îmbunătățirea CV-ului în vederea acestei poziții
        4. Procentajul general de potrivire (evaluat pe o scară de la 0 la 100%)
        5. Fiți cât mai concret, citați cerințele specifice din descrierea postului și punctele din CV.
        """

        try:
            with st.spinner("Generăm o analiză detaliată…"):
                data = {
                    "model": "deepseek-chat",
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.3
                }

                response = requests.post(url, headers=headers, json=data)
                response.raise_for_status()

                result = response.json()
                analysis = result['choices'][0]['message']['content']

                # Сохраняем анализ в сессию, чтобы не пропадал
                st.session_state.analysis = analysis

        except Exception as e:
            st.error(f"Ошибка при запросе к API: {str(e)}")

# Показываем анализ, если он уже есть в сессии
if 'analysis' in st.session_state and st.session_state.analysis:
    st.markdown("## 📊 Analiză detaliată a conformității")
    st.markdown(st.session_state.analysis)

    # Кнопка для скачивания анализа в docx
    def create_word_document(text):
        doc = Document()
        doc.add_heading('Analiză detaliată a conformității', 0)
        for line in text.split('\n'):
            if line.strip():
                if line.startswith('##'):
                    doc.add_heading(line.replace('##', '').strip(), level=1)
                elif line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), level=2)
                else:
                    doc.add_paragraph(line)
        return doc

    doc = create_word_document(st.session_state.analysis)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    st.download_button(
        label="💾 Descarcă analiza (DOCX)",
        data=bio,
        file_name="analiza_potrivire.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.info("Apăsați butonul pentru a genera analiza de potrivire.")

# --- Интервью ---

def generate_interview_questions(cv_text):
    prompt = f"""
    Generează 10 întrebări pentru un interviu introductiv pe baza acestui CV:
    {cv_text}

    Cerințe:

    1. 3 întrebări despre experiența profesională
    2. 2 întrebări despre abilitățile tehnice
    3. 1 întrebare despre punctele slabe
    4. 1 întrebare despre motivație
    5. 1 întrebare despre așteptările salariale
    6. 2 întrebări biografice

    Întrebările trebuie să fie specifice și legate de CV

    Returnează doar o listă numerotată de întrebări, fără explicații suplimentare.
    """

    response = requests.post(
        url,
        headers=headers,
        json={
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3
        }
    )
    return response.json()['choices'][0]['message']['content']

def generate_candidate_profile(questions, answers):
    prompt = f"""
    Pe baza acestor întrebări și răspunsuri, creează un profil al candidatului:

    Întrebări:
    {questions}

    Răspunsuri:

    {answers}

    Структура профиля:
    ### 🧑‍💻 Portret profesional
    - Competențe principale
    - Experiență relevantă
    - Expertiză tehnică

    ### 🎯 Motivație și obiective
    - Interese profesionale
    - Așteptări de la job

    ### 📈 Puncte forte
    - Avantaje cheie
    - Competențe unice

    ### ⚠️ Zone de dezvoltare
    - Puncte slabe
    - Competențe de îmbunătățit

    ### 💰 Așteptări privind compensația
    - Așteptări salariale
    - Disponibilitate pentru negociere
    """

    response = requests.post(
        url,
        headers=headers,
        json={
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2
        }
    )
    return response.json()['choices'][0]['message']['content']

st.title("🤖 AI HR-Recruiter: Interviu introductiv")

# Инициализация состояния интервью
if 'interview_started' not in st.session_state:
    st.session_state.interview_started = False
    st.session_state.questions = None
    st.session_state.answers = {}
    st.session_state.profile = None

if not st.session_state.interview_started:
    if st.button("🎤 A trece interviul introductiv", type="primary"):
        with st.spinner("Pregătim întrebările..."):
            st.session_state.questions = generate_interview_questions(documents[0])
            st.session_state.interview_started = True
        st.rerun()
        

if st.session_state.interview_started:
    st.success("Interviul a început! Vă rog să răspundeți la întrebările de mai jos.")

    questions_list = [q for q in st.session_state.questions.split('\n') if q.strip()]
    for i, question in enumerate(questions_list[:10]):
        st.session_state.answers[i] = st.text_area(
            label=f"**{i+1}:** {question}",
            value=st.session_state.answers.get(i, ""),
            key=f"answer_{i}"
        )

    ########################################################################################

    if st.button("✅ Interviul s-a încheiat", type="primary"):
    
        with st.spinner("Analizăm răspunsurile..."):
            questions_list = [q for q in st.session_state.questions.split('\n') if q.strip()]
            
            formatted_answers = "\n".join(
                [
                    f"{i+1}. {q}\n   Ответ: {st.session_state.answers.get(i, '').strip() or 'Candidatul nu a putut răspunde la această întrebare'}"
                    for i, q in enumerate(questions_list[:10])
                ]
            )
    
            # 🧠 Проверка на ИИ-сгенерированные ответы
            suspicious_flags = []
            for i, q in enumerate(questions_list[:10]):
                answer = st.session_state.answers.get(i, '').strip()
                verdict = check_if_ai_generated(answer)  # Используй свою LLM-функцию
                if 'ai' in verdict.lower():
                    suspicious_flags.append((i+1, verdict))
    
            if suspicious_flags:
                st.warning("🚨 Unele răspunsuri par a fi generate de AI:")
                for q_num, reason in suspicious_flags:
                    st.markdown(f"- Întrebarea {q_num}: răspuns suspectat ca fiind generat de AI")
    
            # 🧾 Генерация профиля кандидата
            st.session_state.profile = generate_candidate_profile(
                st.session_state.questions,
                formatted_answers
            )
    
        st.success("Interviul s-a încheiat!")
        st.balloons()
        st.rerun()


    

      

if st.session_state.profile:
    st.markdown("## 📌 Profilul candidatului")
    st.markdown(st.session_state.profile)

    # Создание и скачивание DOCX профиля кандидата
    def create_word_document_profile(profile_text):
        doc = Document()
        doc.add_heading('Profil Candidat', 0)
        for line in profile_text.split('\n'):
            if line.strip():
                if line.startswith('###'):
                    doc.add_heading(line.replace('###', '').strip(), level=2)
                else:
                    doc.add_paragraph(line)
        return doc

    doc_profile = create_word_document_profile(st.session_state.profile)
    bio_profile = io.BytesIO()
    doc_profile.save(bio_profile)
    bio_profile.seek(0)

    st.download_button(
        label="💾 Descarcă profilul candidatului (DOCX)",
        data=bio_profile,
        file_name="profil_candidat.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


###########################################################################################################
# --- Технические вопросы (генерация через LLM) ---
def generate_technical_questions(cv_text):
    prompt = f"""
    Generează 5 întrebări tehnice specifice pe baza acestui CV:
    {cv_text}

    Întrebările trebuie să testeze competențele tehnice ale candidatului.
    Returnează doar o listă numerotată de întrebări, fără alte explicații.
    """
    response = requests.post(
        url,
        headers=headers,
        json={
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3
        }
    )
    return response.json()['choices'][0]['message']['content']

# --- Генерация технического фидбека по ответам ---
def generate_technical_feedback(questions, answers):
    prompt = f"""
    Pe baza următoarelor întrebări tehnice și răspunsuri, oferă un feedback detaliat și un scor evaluativ (0-10) pentru competențele tehnice ale candidatului.

    Întrebări:
    {questions}

    Răspunsuri:
    {answers}

    Formatează răspunsul astfel:

    Feedback detaliat:
    [text]

    Scor tehnic: [număr de la 0 la 10]
    """
    response = requests.post(
        url,
        headers=headers,
        json={
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2
        }
    )
    return response.json()['choices'][0]['message']['content']

# --- Итоговая рекомендация ---
def generate_final_recommendation(profile, tech_feedback, ai_flags_count):
    prompt = f"""
    Având următorul profil al candidatului:

    {profile}

    Feedback tehnic:

    {tech_feedback}

    Număr de răspunsuri suspectate ca fiind generate de AI: {ai_flags_count}

    Pe baza acestor informații, formulează o concluzie finală clară cu una din următoarele recomandări:
    - Recomandare pentru angajare
    - Recomandare cu rezerve
    - Refuz argumentat

    Include argumentele principale pentru decizie.
    """
    response = requests.post(
        url,
        headers=headers,
        json={
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2
        }
    )
    return response.json()['choices'][0]['message']['content']

# --- Streamlit UI ---

st.title("🤖 AI HR-Recruiter: Interviu tehnic și concluzie finală")

if 'tech_interview_started' not in st.session_state:
    st.session_state.tech_interview_started = False
    st.session_state.tech_questions = None
    st.session_state.tech_answers = {}
    st.session_state.tech_feedback = None
    st.session_state.final_recommendation = None

# Кнопка начала технического интервью
if not st.session_state.tech_interview_started and st.session_state.interview_started:
    if st.button("💻 Începe interviul tehnic"):
        with st.spinner("Pregătim întrebările tehnice..."):
            st.session_state.tech_questions = generate_technical_questions(documents[0])
            st.session_state.tech_interview_started = True
        st.rerun()

# Форма технического интервью
if st.session_state.tech_interview_started:
    st.success("Interviul tehnic a început! Vă rugăm să răspundeți la întrebările de mai jos.")

    tech_q_list = [q for q in st.session_state.tech_questions.split('\n') if q.strip()]
    for i, question in enumerate(tech_q_list[:5]):
        st.session_state.tech_answers[i] = st.text_area(
            label=f"**{i+1}:** {question}",
            value=st.session_state.tech_answers.get(i, ""),
            key=f"tech_answer_{i}"
        )

    if st.button("✅ Finalizează interviul tehnic"):
        with st.spinner("Analizăm răspunsurile tehnice..."):
            formatted_tech_answers = "\n".join(
                [
                    f"{i+1}. {q}\n   Răspuns: {st.session_state.tech_answers.get(i, '').strip() or 'Nu a răspuns'}"
                    for i, q in enumerate(tech_q_list[:5])
                ]
            )

            # Анализ технических ответов через LLM
            st.session_state.tech_feedback = generate_technical_feedback(
                st.session_state.tech_questions,
                formatted_tech_answers
            )

            # Проверка AI-сгенерированных ответов из всех этапов (здесь вызываем свою функцию)
            suspicious_flags = []
            all_answers = list(st.session_state.answers.values()) + list(st.session_state.tech_answers.values())
            for idx, ans in enumerate(all_answers):
                verdict = check_if_ai_generated(ans)
                if 'ai' in verdict.lower():
                    suspicious_flags.append((idx+1, verdict))

            # Генерация итоговой рекомендации
            st.session_state.final_recommendation = generate_final_recommendation(
                st.session_state.profile,
                st.session_state.tech_feedback,
                len(suspicious_flags)
            )
        st.success("Interviul tehnic s-a încheiat!")
        st.balloons()
        st.rerun()




# --- После завершения технического интервью ---

# 1. Сначала выводим технический фидбек (если есть)
if st.session_state.tech_feedback:
    st.markdown("## 💻 Feedback tehnic")
    st.markdown(st.session_state.tech_feedback)

# 2. Затем выводим финальную рекомендацию (если есть)
if st.session_state.final_recommendation:
    st.markdown("## 📋 Concluzia finală")
    st.markdown(st.session_state.final_recommendation)
    
    # Создаем полный отчёт для скачивания (включая фидбек и рекомендацию)
    def create_final_report():
        doc = Document()
        doc.add_heading("Raport complet candidat", 0)
        
        # Профиль
        doc.add_heading("Profil candidat", 1)
        for line in st.session_state.profile.split('\n'):
            if line.strip():
                if line.startswith('###'):
                    doc.add_heading(line.replace('###', '').strip(), 2)
                else:
                    doc.add_paragraph(line)
        
        # Технический фидбек
        doc.add_heading("Feedback tehnic", 1)
        doc.add_paragraph(st.session_state.tech_feedback)
        
        # Рекомендация
        doc.add_heading("Recomandare finală", 1)
        doc.add_paragraph(st.session_state.final_recommendation)
        
        return doc

    # Кнопка скачивания
    report_doc = create_final_report()
    report_bio = io.BytesIO()
    report_doc.save(report_bio)
    report_bio.seek(0)
    
    st.download_button(
        label="💾 Descarcă raport complet (DOCX)",
        data=report_bio,
        file_name="raport_candidat.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Кнопка сброса
    if st.button("🔄 Resetează procesul"):
        for key in st.session_state.keys():
            del st.session_state[key]
        st.rerun()
